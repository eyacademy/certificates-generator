import io
import os
import csv
import re
import zipfile
import tempfile
import subprocess
import logging
import shutil
from typing import Dict, List, Optional, Tuple

try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except Exception:
    HAS_XLSX = False

from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import (
    FileResponse,
    HTMLResponse,
    Response,
    StreamingResponse,
    PlainTextResponse,
)
from fastapi.middleware.cors import CORSMiddleware

import asyncio
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
import json
import time

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docxtpl import DocxTemplate


# -----------------------------------------------------------------------------
# App / logging
# -----------------------------------------------------------------------------
app = FastAPI(title="Certificates Generator")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("certefikati")

# === Online visual tweaks (padding with non-breaking spaces) ===
ONLINE_NAME_PAD_NBSP = 2    # сколько \u00A0 добавить к Имени в online
ONLINE_COURSE_PAD_NBSP = 2  # сколько \u00A0 добавить к Тренингу в online (чуть дальше вправо)
NBSP = "\u00A0"


@app.head("/")
def head_root():
    return Response(status_code=200)


@app.get("/")
def root():
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url="/ui")


@app.get("/ui")
def ui():
    html_content = """
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Генератор сертификатов</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            max-width: 600px;
            margin: 50px auto;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1 { color: #333; text-align: center; margin-bottom: 30px; }
        .form-group { margin-bottom: 20px; }
        label { display: block; margin-bottom: 5px; font-weight: bold; color: #555; }
        input[type="file"] {
            width: 100%; padding: 10px; border: 2px dashed #ddd; border-radius: 5px; background: #fafafa;
        }
        .radio-group { display: flex; gap: 20px; margin-top: 10px; }
        .radio-item { display: flex; align-items: center; gap: 5px; }
        input[type="radio"] { margin: 0; }
        button {
            background: #007bff; color: white; border: none; padding: 12px 30px; border-radius: 5px;
            cursor: pointer; font-size: 16px; width: 100%; margin-top: 20px;
        }
        button:hover { background: #0056b3; }
        button:disabled { background: #ccc; cursor: not-allowed; }
        .status { margin-top: 20px; padding: 10px; border-radius: 5px; display: none; }
        .status.success { background: #d4edda; color: #155724; border: 1px solid #c3e6cb; }
        .status.error { background: #f8d7da; color: #721c24; border: 1px solid #f5c6cb; }
        .progress { width: 100%; height: 20px; background: #f0f0f0; border-radius: 10px; overflow: hidden; margin-top: 10px; display: none; }
        .progress-bar { height: 100%; background: #007bff; width: 0%; transition: width 0.3s; }
        .progress-info { margin-top: 6px; color: #555; font-size: 14px; display: none; }
    </style>
</head>
<body>
    <div class="container">
        <h1>Генератор сертификатов</h1>

        <form id="certificateForm">
            <div class="form-group">
                <label for="csvFile">Выберите файл (CSV или Excel):</label>
                <input type="file" id="csvFile" name="csv_file" accept=".csv,.xlsx,.xls" required>
                <div id="fileStatus">Файл не выбран</div>
            </div>

            <div class="form-group">
                <label>Тип сертификата:</label>
                <div class="radio-group">
                    <div class="radio-item">
                        <input type="radio" id="print" name="mode" value="print" checked>
                        <label for="print">Печать</label>
                    </div>
                    <div class="radio-item">
                        <input type="radio" id="online" name="mode" value="online">
                        <label for="online">Онлайн</label>
                    </div>
                </div>
            </div>

            <button type="submit" id="generateBtn">Сгенерировать</button>
        </form>

        <div class="progress" id="progress"><div class="progress-bar" id="progressBar"></div></div>
        <div class="progress-info" id="progressInfo"></div>
        <div class="status" id="status"></div>
    </div>

    <script>
        const form = document.getElementById('certificateForm');
        const fileInput = document.getElementById('csvFile');
        const fileStatus = document.getElementById('fileStatus');
        const generateBtn = document.getElementById('generateBtn');
        const progress = document.getElementById('progress');
        const progressBar = document.getElementById('progressBar');
        const progressInfo = document.getElementById('progressInfo');
        const status = document.getElementById('status');

        fileInput.addEventListener('change', function() {
            if (this.files.length > 0) {
                fileStatus.textContent = `Выбран файл: ${this.files[0].name}`;
                fileStatus.style.color = '#28a745';
            } else {
                fileStatus.textContent = 'Файл не выбран';
                fileStatus.style.color = '#dc3545';
            }
        });

        form.addEventListener('submit', async function(e) {
            e.preventDefault();

            const formData = new FormData();
            const file = fileInput.files[0];
            const mode = document.querySelector('input[name="mode"]:checked').value;

            if (!file) {
                showStatus('Пожалуйста, выберите CSV файл', 'error');
                return;
            }

            formData.append('csv_file', file);
            formData.append('mode', mode);

            const jobId = (window.crypto && crypto.randomUUID)
                ? crypto.randomUUID() : ('job-' + Date.now() + '-' + Math.random().toString(16).slice(2));
            formData.append('job_id', jobId);

            generateBtn.disabled = true;
            generateBtn.textContent = 'Генерация...';
            progress.style.display = 'block';
            progressInfo.style.display = 'block';
            status.style.display = 'none';

            // progress via SSE
            const es = new EventSource(`/progress/${jobId}`);
            let sseStage = 'init';
            es.onmessage = (ev) => {
                try {
                    const data = JSON.parse(ev.data || '{}');
                    sseStage = data.stage || sseStage;
                    if (sseStage === 'processing' || sseStage === 'zipping') {
                        updateProgress(data.percent || 0);
                        setInfo(`${data.message || ''} (${data.processed || 0}/${data.total || 0})`);
                    } else if (sseStage === 'done') {
                        setInfo('Подготовка к скачиванию...');
                        downloadZip(jobId);
                    } else if (sseStage === 'uploading') {
                        setInfo('Загрузка файла...');
                    } else if (sseStage === 'error') {
                        setInfo(data.message || 'Ошибка');
                    }
                    if (sseStage === 'done' || sseStage === 'error') es.close();
                } catch (e) {}
            };
            es.onerror = () => es.close();

            try {
                const xhr = new XMLHttpRequest();
                xhr.open('POST', '/generate-async');
                xhr.responseType = 'json';
                xhr.upload.onprogress = (e) => {
                    if (e.lengthComputable) {
                        const pct = Math.round((e.loaded / e.total) * 100);
                        updateProgress(Math.min(pct, 99));
                        setInfo(`Загрузка файла: ${pct}%`);
                    } else {
                        setInfo('Загрузка файла...');
                    }
                };
                xhr.onload = () => {
                    if (xhr.status !== 200) {
                        const text = typeof xhr.response === 'string' ? xhr.response : (xhr.response?.detail || 'Ошибка запуска');
                        showStatus(`Ошибка: ${text}`, 'error');
                        cleanup();
                        return;
                    }
                    setInfo('Файл загружен, идёт обработка...');
                };
                xhr.onerror = () => {
                    showStatus('Сетевая ошибка при запросе', 'error');
                    cleanup();
                };
                xhr.send(formData);
            } catch (error) {
                showStatus(`Ошибка сети: ${error.message}`, 'error');
                cleanup();
            }
        });

        function showStatus(message, type) {
            status.textContent = message;
            status.className = `status ${type}`;
            status.style.display = 'block';
        }
        function updateProgress(pct) { progressBar.style.width = (pct || 0) + '%'; }
        function setInfo(text) { progressInfo.textContent = text || ''; }
        function cleanup() {
            generateBtn.disabled = false;
            generateBtn.textContent = 'Сгенерировать';
            setTimeout(() => {
                progress.style.display = 'none';
                progressInfo.style.display = 'none';
                progressBar.style.width = '0%';
                progressInfo.textContent = '';
            }, 800);
        }
        async function downloadZip(jobId) {
            try {
                const res = await fetch(`/download/${jobId}`);
                if (!res.ok) { setTimeout(() => downloadZip(jobId), 1500); return; }
                const blob = await res.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url; a.download = 'certificates.zip';
                document.body.appendChild(a); a.click();
                document.body.removeChild(a); window.URL.revokeObjectURL(url);
                updateProgress(100); setInfo('Готово');
                showStatus('Сертификаты успешно сгенерированы!', 'success');
                cleanup();
            } catch (e) { showStatus('Ошибка при скачивании результата', 'error'); cleanup(); }
        }
    </script>
</body>
</html>
    """
    return HTMLResponse(content=html_content)


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# -----------------------------------------------------------------------------
# Paths / fonts
# -----------------------------------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "Templates"))
PDF_TEMPLATES_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "TemplatesPDF"))

FONTS_DIR_CANDIDATES = [
    os.path.abspath(os.path.join(BASE_DIR, "..", "fonts")),
    os.path.abspath(os.path.join(BASE_DIR, "..", "@fonts")),
]
FONTS_DIR = next((p for p in FONTS_DIR_CANDIDATES if os.path.isdir(p)), FONTS_DIR_CANDIDATES[0])

FONT_NAME = "EYInterstate"
registered = False
try:
    regular = os.path.join(FONTS_DIR, "EYINTERSTATE-REGULAR.OTF")
    if os.path.exists(regular):
        pdfmetrics.registerFont(TTFont(FONT_NAME, regular))
        registered = True
        for (alias, path) in [
            ("EYInterstate-Bold", os.path.join(FONTS_DIR, "EYINTERSTATE-BOLD.OTF")),
            ("EYInterstate-Light", os.path.join(FONTS_DIR, "EYINTERSTATE-LIGHT.OTF")),
        ]:
            if os.path.exists(path):
                try: pdfmetrics.registerFont(TTFont(alias, path))
                except Exception: pass
except Exception:
    registered = False

if not registered:
    FONT_NAME = "DejaVuSans"
    for p in ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
              "/usr/share/fonts/trruetype/dejavu/DejaVuSansCondensed.ttf",
              "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf"]:
        if os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont(FONT_NAME, p))
                registered = True
                break
            except Exception:
                continue
if not registered:
    FONT_NAME = "Helvetica"


# -----------------------------------------------------------------------------
# DOCX map
# -----------------------------------------------------------------------------
DOCX_MAP: Dict[str, Dict[str, Dict[str, str]]] = {
    "print": {
        "duration_day": {"normal": "template_duration_day.docx", "small": "template_small_duration_day.docx"},
        "2day_2month": {"normal": "template2_2day_2month.docx", "small": "template2_small_2day_2month.docx"},
        "1day_1month": {"normal": "template3_1day_1month.docx", "small": "template3_small_1day_1month.docx"},
    },
    "online": {
        "duration_day": {"normal": "template-online_duration_day.docx", "small": "template-online_small_duration_day.docx"},
        "2day_2month": {"normal": "template-online2_2day_2month.docx", "small": "template-online2_small_2day_2month.docx"},
        "1day_1month": {"normal": "template-online3_1day_1month.docx", "small": "template-online3_small_1day_1month.docx"},
    },
}

DOCX_TO_PDF_CACHE: Dict[str, str] = {}


# -----------------------------------------------------------------------------
# Helpers: CSV/Excel parsing, dates, name sizing
# -----------------------------------------------------------------------------
MONTH_GEN = {
    1: "January", 2: "February", 3: "March", 4: "April", 5: "May", 6: "June",
    7: "July", 8: "August", 9: "September", 10: "October", 11: "November", 12: "December",
}

def _norm_key(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = s.replace("ё", "е")
    return s

KEY_ALIASES = {
    "first_name": {"имя", "name", "first name", "first_name", "имя/name", "name/имя"},
    "last_name":  {"фамилия", "surname", "last name", "last_name", "фамилия/surname", "surname/фамилия"},
    "course":     {"название тренинга", "название", "course", "course name", "course_name",
                   "название тренинга/название", "название/название тренинга"},
    "dates":      {"даты", "дата", "dates", "date", "даты/дата", "date/dates"},
    "id":         {"id", "id/id", "идентификатор", "certificate id", "сертификат id"},
    "city":       {"город", "city", "город/city", "city/город"},
    "country":    {"страна", "country", "страна/country", "country/страна"},
}

def _build_row_with_normalized_keys(row: Dict[str, str]) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for k, v in (row or {}).items():
        if k is None: continue
        val = (v or "").strip()
        out[k] = val
        nk = _norm_key(str(k))
        out[nk] = val
        if "/" in str(k):
            parts = [p.strip() for p in str(k).split("/") if p.strip()]
            for p in parts:
                out[p] = val
                out[_norm_key(p)] = val
    return out

def _get_field(row: Dict[str, str], canonical: str) -> str:
    r = _build_row_with_normalized_keys(row)
    aliases = KEY_ALIASES.get(canonical, set())
    for key in list(r.keys()):
        nk = _norm_key(str(key))
        if nk in aliases: return (r[key] or "").strip()
    return (r.get(canonical) or r.get(_norm_key(canonical)) or "").strip()

def detect_delimiter(text: str) -> str:
    sample = text[:4096]
    candidates = [',', ';', '\t']
    counts = {c: sample.count(c) for c in candidates}
    delim = max(counts, key=counts.get)
    return delim if counts[delim] > 0 else ','

def normalize_csv_and_get_delimiter(text: str) -> Tuple[str, str]:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    delim = detect_delimiter(text)
    reader = csv.reader(io.StringIO(text), delimiter=delim)
    rows = [list(map(lambda s: (s or '').strip(), r)) for r in reader]
    rows = [r for r in rows if any(cell != '' for cell in r)]
    if not rows: return '', delim

    def is_generic_header(cells: List[str]) -> bool:
        return bool(cells) and all(cell.lower().startswith('column') or cell.lower().startswith('unnamed') for cell in cells)

    start_idx = 1
    header = rows[0]
    if is_generic_header(header) and len(rows) >= 2:
        start_idx = 2
        header = rows[1]
    if header:
        header[0] = header[0].lstrip('\ufeff')
    norm_rows = [header] + rows[start_idx:]
    normalized_text = '\n'.join(delim.join(r) for r in norm_rows)
    return normalized_text, delim

def parse_dates(s: str) -> Dict[str, Optional[int]]:
    s = (s or "").strip()
    date_pattern = r"\b(\d{1,2})\.(\d{1,2})\.(\d{2,4})\b"
    dates = list(re.finditer(date_pattern, s))
    if len(dates) >= 2:
        d1, m1, y1 = int(dates[0].group(1)), int(dates[0].group(2)), int(dates[0].group(3))
        d2, m2, y2 = int(dates[1].group(1)), int(dates[1].group(2)), int(dates[1].group(3))
        if y1 < 100: y1 += 2000
        if y2 < 100: y2 += 2000
        return {"d1": d1, "m1": m1, "d2": d2, "m2": m2, "y": y1}
    elif len(dates) == 1:
        d1, m1, y1 = int(dates[0].group(1)), int(dates[0].group(2)), int(dates[0].group(3))
        if y1 < 100: y1 += 2000
        return {"d1": d1, "m1": m1, "d2": None, "m2": None, "y": y1}

    s_lower = s.lower()
    m_year = re.search(r"\b(20\d{2})\b", s)
    y = int(m_year.group(1)) if m_year else __import__("datetime").datetime.now().year

    months_tokens = [
        ("January", 1), ("February", 2), ("March", 3), ("April", 4),
        ("May", 5), ("June", 6), ("July", 7), ("August", 8),
        ("September", 9), ("October", 10), ("November", 11), ("December", 12),
    ]
    def detect_month(token: str) -> Optional[int]:
        if token.isdigit() and 1 <= int(token) <= 12: return int(token)
        for pref, m in months_tokens:
            if token.startswith(pref.lower()): return m
        return None

    tokens = re.sub(r"[,;]+", " ", s_lower)
    tokens = re.sub(r"\s+", " ", tokens).strip().split(" ")
    days = [int(t) for t in tokens if t.isdigit() and 1 <= int(t) <= 31]
    month_idx = [(i, detect_month(t)) for i, t in enumerate(tokens) if detect_month(t)]

    rng = re.search(r"\b(\d{1,2})\s*[-–]\s*(\d{1,2})\b", s_lower)
    if rng and month_idx:
        return {"d1": int(rng.group(1)), "m1": month_idx[0][1],
                "d2": int(rng.group(2)), "m2": month_idx[0][1], "y": y}
    if len(days) >= 2 and len(month_idx) >= 2:
        return {"d1": days[0], "m1": month_idx[0][1], "d2": days[1], "m2": month_idx[1][1], "y": y}
    if len(days) >= 2 and len(month_idx) == 1:
        return {"d1": days[0], "m1": month_idx[0][1], "d2": days[1], "m2": month_idx[0][1], "y": y}
    if len(days) >= 1 and len(month_idx) >= 1:
        return {"d1": days[0], "m1": month_idx[0][1], "d2": None, "m2": None, "y": y}
    return {"d1": days[0] if days else 1, "m1": month_idx[0][1] if month_idx else 1, "d2": None, "m2": None, "y": y}

def format_dates_for_jinja(p: Dict[str, Optional[int]]) -> Dict[str, str]:
    d1, m1, d2, m2, y = p["d1"], p["m1"], p["d2"], p["m2"], p["y"]
    result = {
        "Имя": "", "Фамилия": "", "Тренинг": "",
        "Год": str(y), "Город": "Москва",
    }
    if d1 and d2:
        result.update({"Дата1": str(d1), "Дата2": str(d2),
                       "Месяц1": MONTH_GEN[m1], "Месяц2": MONTH_GEN[m2]})
    else:
        result.update({"Дата1": str(d1), "Месяц1": MONTH_GEN[m1]})
    return result

def pick_kind(p: Dict[str, Optional[int]]) -> str:
    d1, m1, d2, m2 = p["d1"], p["m1"], p["d2"], p["m2"]
    if d1 and d2:
        return "duration_day" if m1 == m2 else "2day_2month"
    return "1day_1month"

def string_width_pt(text: str, size: int) -> float:
    try:
        return pdfmetrics.stringWidth(text, FONT_NAME, size)
    except Exception:
        return size * max(1, len(text)) * 0.5

def need_small_variant(full_name: str) -> bool:
    base_size = 28
    max_width = 400
    w = string_width_pt(full_name, base_size)
    return w > max_width

def sanitize_filename(s: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", s).replace(" ", "_")[:100]


# -----------------------------------------------------------------------------
# DOCX -> PDF (LibreOffice)
# -----------------------------------------------------------------------------
def docx_to_pdf_cached(docx_path: str) -> str:
    abs_docx = os.path.abspath(docx_path)
    if abs_docx in DOCX_TO_PDF_CACHE:
        logger.info(f"Using cached PDF for {abs_docx}")
        return DOCX_TO_PDF_CACHE[abs_docx]

    logger.info(f"Converting DOCX to PDF: {abs_docx}")
    out_dir = tempfile.mkdtemp(prefix="docx2pdf_")

    libreoffice_paths = [
        "soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\LibreOffice\program\soffice.exe",
    ]
    cmd = None
    for path in libreoffice_paths:
        try:
            if path == "soffice":
                test_cmd = ["soffice", "--version"]
                subprocess.run(test_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
                cmd = [path, "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_docx]
                break
            elif os.path.exists(path):
                cmd = [path, "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_docx]
                break
        except (subprocess.CalledProcessError, FileNotFoundError):
            continue
    if cmd is None:
        raise RuntimeError("LibreOffice not found. Install it or add to PATH.")

    profile_dir = os.path.join(out_dir, "lo_profile")
    try: os.makedirs(profile_dir, exist_ok=True)
    except Exception: pass
    profile_url = "file:///" + os.path.abspath(profile_dir).replace("\\", "/").lstrip("/")

    cmd_with_profile = [
        cmd[0], "--headless", "--norestore", "--nolockcheck",
        f"-env:UserInstallation={profile_url}",
        "--convert-to", "pdf", "--outdir", out_dir, abs_docx,
    ]
    proc = subprocess.run(cmd_with_profile, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

    pdf_path = os.path.join(out_dir, os.path.splitext(os.path.basename(abs_docx))[0] + ".pdf")
    if not os.path.exists(pdf_path):
        stderr_txt = proc.stderr.decode(errors='ignore') if proc.stderr else ''
        stdout_txt = proc.stdout.decode(errors='ignore') if proc.stdout else ''
        raise RuntimeError(f"LibreOffice convert failed: {stderr_txt or stdout_txt or 'unknown error'}")

    DOCX_TO_PDF_CACHE[abs_docx] = pdf_path
    return pdf_path


# -----------------------------------------------------------------------------
# Render DOCX + (для online) поправка отступа курса
# -----------------------------------------------------------------------------
def render_docx_template(
    docx_path: str,
    context: Dict[str, str],
    adjust_online_course_indent: bool = False,
    course_indent_pts: int = 18,   # запасной отступ для обычных параграфов
) -> bytes:
    """Рендерит DOCX и (в online-режиме) сдвигает курс вправо так,
    чтобы переносы строк в текстбоксе тоже были смещены."""
    from docxtpl import DocxTemplate
    import zipfile
    from xml.etree import ElementTree as ET

    # 1) Рендер шаблона во временный DOCX
    doc = DocxTemplate(docx_path)
    doc.render(context)
    tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp_docx.name)
    tmp_docx.close()

    if adjust_online_course_indent:
        # 2) Попытка сдвигать обычные параграфы через python-docx (если курс НЕ в текстбоксе)
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            d = Document(tmp_docx.name)

            def all_paragraphs(docx):
                pars = list(docx.paragraphs)
                def walk_tables(tables, acc):
                    for t in tables:
                        for row in t.rows:
                            for cell in row.cells:
                                acc.extend(cell.paragraphs)
                                if cell.tables:
                                    walk_tables(cell.tables, acc)
                walk_tables(docx.tables, pars)
                return pars

            def norm(s: str) -> str:
                return " ".join((s or "").split()).lower()

            paragraphs = all_paragraphs(d)
            target_course = norm((context or {}).get("Тренинг", ""))
            target_name   = norm((context or {}).get("Имя", ""))

            name_indent = None
            name_align = None
            if target_name:
                for p in paragraphs:
                    if target_name[:20] in norm(p.text):
                        name_indent = p.paragraph_format.left_indent
                        name_align = p.alignment
                        break

            # найдём курс среди обычных параграфов
            course_para = None
            if target_course:
                head = " ".join(target_course.split()[:3])
                for p in paragraphs:
                    if head and head in norm(p.text):
                        course_para = p
                        break

            if course_para is not None:
                # выставим отступ абзацу
                if name_indent is not None:
                    course_para.paragraph_format.left_indent = name_indent
                else:
                    course_para.paragraph_format.left_indent = Pt(course_indent_pts)
                course_para.paragraph_format.first_line_indent = Pt(0)
                if name_align in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
                    course_para.alignment = name_align
                d.save(tmp_docx.name)
        except Exception as e:
            logging.warning(f"Docx paragraph adjust skipped: {e}")

        # 3) ГЛАВНОЕ: сдвигаем абзац курса внутри ТЕКСТБОКСОВ (wps:txbx и v:textbox)
        try:
            ns = {
                "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
                "v":   "urn:schemas-microsoft-com:vml",
            }
            # сколько сдвигать в текстбоксе (twips) — ~0.5 см
            LEFT_TWIPS = 360

            def norm(s: str) -> str:
                return " ".join((s or "").split()).lower()

            target_course = norm((context or {}).get("Тренинг", ""))
            head = " ".join(target_course.split()[:3]) if target_course else ""
            if head:
                with zipfile.ZipFile(tmp_docx.name, "a") as z:
                    # какие части патчим
                    candidates = ["word/document.xml"] + \
                                 [n for n in z.namelist() if n.startswith("word/header") and n.endswith(".xml")]

                    for part in candidates:
                        try:
                            xml = z.read(part)
                        except KeyError:
                            continue
                        root = ET.fromstring(xml)

                        def para_text(p):
                            return norm("".join(t.text or "" for t in p.findall(".//w:t", ns)))

                        changed = False
                        # параграфы внутри новых текстбоксов
                        for p in root.findall(".//wps:txbx//w:p", ns):
                            if head in para_text(p):
                                pPr = p.find("w:pPr", ns) or ET.SubElement(p, "{%s}pPr" % ns["w"])
                                ind = pPr.find("w:ind", ns) or ET.SubElement(pPr, "{%s}ind" % ns["w"])
                                ind.set("{%s}left" % ns["w"], str(LEFT_TWIPS))
                                ind.set("{%s}firstLine" % ns["w"], "0")
                                changed = True
                        # параграфы внутри старых VML-текстбоксов
                        for p in root.findall(".//v:textbox//w:p", ns):
                            if head in para_text(p):
                                pPr = p.find("w:pPr", ns) or ET.SubElement(p, "{%s}pPr" % ns["w"])
                                ind = pPr.find("w:ind", ns) or ET.SubElement(pPr, "{%s}ind" % ns["w"])
                                ind.set("{%s}left" % ns["w"], str(LEFT_TWIPS))
                                ind.set("{%s}firstLine" % ns["w"], "0")
                                changed = True

                        if changed:
                            new_xml = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                            z.writestr(part, new_xml)
        except Exception as e:
            logging.warning(f"Textbox indent adjust skipped: {e}")

    # 4) Конвертация в PDF
    pdf_path = docx_to_pdf_cached(tmp_docx.name)
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()
    os.unlink(tmp_docx.name)
    return pdf_bytes


# -----------------------------------------------------------------------------
# SSE progress
# -----------------------------------------------------------------------------
@dataclass
class ProgressState:
    total: int = 0
    processed: int = 0
    stage: str = "init"  # init | uploading | processing | zipping | done | error
    message: str = ""
    errors: int = 0
    created: float = field(default_factory=lambda: time.time())
    queue: asyncio.Queue = field(default_factory=asyncio.Queue)

PROGRESS: Dict[str, ProgressState] = {}
JOB_RESULTS: Dict[str, bytes] = {}

def get_progress(job_id: str) -> ProgressState:
    if job_id not in PROGRESS:
        PROGRESS[job_id] = ProgressState()
    return PROGRESS[job_id]

def snapshot(state: ProgressState) -> Dict[str, object]:
    percent = int(state.processed * 100 / max(1, state.total)) if state.total > 0 else 0
    return {
        "total": state.total,
        "processed": state.processed,
        "percent": percent,
        "stage": state.stage,
        "message": state.message,
        "errors": state.errors,
    }

async def emit(job_id: str):
    state = get_progress(job_id)
    await state.queue.put("update")

@app.get("/progress/{job_id}")
async def progress_stream(job_id: str):
    async def event_gen():
        state = get_progress(job_id)
        await emit(job_id)
        while True:
            try:
                _ = await asyncio.wait_for(state.queue.get(), timeout=15.0)
            except asyncio.TimeoutError:
                yield "event: ping\ndata: {}\n\n"
                continue
            data = json.dumps(snapshot(state), ensure_ascii=False)
            yield f"data: {data}\n\n"
            if state.stage in ("done", "error"):
                break
    return StreamingResponse(event_gen(), media_type="text/event-stream")


# -----------------------------------------------------------------------------
# Misc endpoints
# -----------------------------------------------------------------------------
@app.get("/health")
def health() -> PlainTextResponse:
    return PlainTextResponse("ok")

@app.get("/sample-excel")
def sample_excel():
    root = os.path.abspath(os.path.join(BASE_DIR, ".."))
    sample_path = os.path.join(root, "Template_Certificates.xlsx")
    if not os.path.exists(sample_path):
        return PlainTextResponse("Template_Certificates.xlsx not found", status_code=404)
    return FileResponse(
        sample_path,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="Template_Certificates.xlsx",
    )

@app.get("/check-templates")
def check_templates():
    missing_templates = []
    available_templates = []
    for group in ["print", "online"]:
        for kind in ["duration_day", "2day_2month", "1day_1month"]:
            for variant in ["normal", "small"]:
                try:
                    docx_name = DOCX_MAP[group][kind][variant]
                    docx_path = os.path.join(TEMPLATES_DIR, docx_name)
                    if os.path.exists(docx_path):
                        available_templates.append(f"{group}/{kind}/{variant}: {docx_name}")
                    else:
                        missing_templates.append(f"{group}/{kind}/{variant}: {docx_name}")
                except KeyError:
                    missing_templates.append(f"{group}/{kind}/{variant}: NOT_FOUND_IN_MAP")
    return {
        "available_templates": available_templates,
        "missing_templates": missing_templates,
        "templates_dir": TEMPLATES_DIR,
        "templates_dir_exists": os.path.exists(TEMPLATES_DIR)
    }


# -----------------------------------------------------------------------------
# Core: generate (sync → returns zip), generate-async (background)
# -----------------------------------------------------------------------------
async def _as_completed_iter(coros):
    for fut in asyncio.as_completed(coros):
        yield fut


@app.post("/generate")
async def generate(
    csv_file: UploadFile = File(...),
    mode: str = Form(...),                  # print | online
    job_id: Optional[str] = Form(None),
):
    try:
        logger.info(f"Starting certificate generation for mode: {mode}")
        state: Optional[ProgressState] = None
        if job_id:
            state = get_progress(job_id)
            state.stage = "uploading"
            state.message = "Загрузка файла"
            await emit(job_id)

        data = await csv_file.read()
        filename = (csv_file.filename or '').lower()
        incoming_fields: List[str] = []
        rows_list: List[Dict[str, str]] = []
        txt = ''
        delim = ','

        if filename.endswith('.xlsx') or filename.endswith('.xlsm'):
            if not HAS_XLSX:
                raise RuntimeError('Поддержка Excel не установлена на сервере')
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            ws = wb.active
            headers: List[str] = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [(c if c is not None else '') for c in row]
                if not any(str(c).strip() for c in cells): continue
                if not headers:
                    candidate = [str(c).strip() for c in cells]
                    tech = all(h.lower().startswith('column') or h.lower().startswith('unnamed') or h == '' for h in candidate)
                    non_empty = [h for h in candidate if h]
                    if tech or len(non_empty) < 3: continue
                    headers = candidate; continue
                d: Dict[str, str] = {}
                for idx, h in enumerate(headers):
                    val = '' if idx >= len(cells) or cells[idx] is None else str(cells[idx])
                    d[h] = val
                rows_list.append(d)
            if not headers or not rows_list:
                raise ValueError('Excel: нераспознан заголовок или нет данных')
        else:
            raw_txt = data.decode('utf-8-sig', errors='ignore')
            txt, delim = normalize_csv_and_get_delimiter(raw_txt)
            if not txt:
                raise ValueError('CSV пустой или нераспознанный формат')
            dict_reader = csv.DictReader(io.StringIO(txt), delimiter=delim)
            incoming_fields = list(dict_reader.fieldnames or [])
            rows_list = [row for row in dict_reader]

        total = len(rows_list)
        if state:
            state.total = total
            state.processed = 0
            state.stage = "processing"
            state.message = "Обработка строк"
            await emit(job_id)

        mem_zip = io.BytesIO()
        processed_count = 0
        loop = asyncio.get_event_loop()
        tasks = []

        with ThreadPoolExecutor(max_workers=2) as executor:
            with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for row_num, row in enumerate(rows_list, 1):
                    try:
                        is_online = (mode == "online")
                        course     = _get_field(row, "course")
                        dates_raw  = _get_field(row, "dates")
                        first_name = _get_field(row, "first_name")
                        last_name  = _get_field(row, "last_name")
                        cert_id    = _get_field(row, "id")
                        city       = _get_field(row, "city")
                        country    = _get_field(row, "country")

                        if not (course and dates_raw and first_name and last_name and cert_id):
                            logger.warning(f"Skipping row {row_num}: missing required fields")
                            continue

                        parsed = parse_dates(dates_raw)
                        kind = pick_kind(parsed)
                        use_small = need_small_variant(f"{first_name} {last_name}")
                        variant = "small" if use_small else "normal"

                        group = "online" if is_online else "print"
                        docx_name = DOCX_MAP[group][kind][variant]
                        docx_path = os.path.join(TEMPLATES_DIR, docx_name)
                        if not os.path.exists(docx_path):
                            raise FileNotFoundError(f"Template not found: {docx_path}")

                        context = format_dates_for_jinja(parsed)
                        context.update({
                            "Имя": first_name,
                            "Фамилия": last_name,
                            "Тренинг": course,
                            "Идентификатор": cert_id,
                            "Город": city or context.get("Город", "Москва"),
                            "Страна": country,
                        })

                        if group == "online":
                            pad_name = NBSP * ONLINE_NAME_PAD_NBSP
                            pad_course = NBSP * ONLINE_COURSE_PAD_NBSP
                            context["Имя"] = pad_name + context.get("Имя", "")
                            context["Тренинг"] = pad_course + context.get("Тренинг", "")

                        async def render_one(docx_path=docx_path, context=context, cert_id=cert_id,
                                             last_name=last_name, first_name=first_name, group=group):
                            adjust = (group == "online")
                            pdf_bytes = await loop.run_in_executor(
                                executor, render_docx_template, docx_path, context, adjust
                            )
                            fname = f"{sanitize_filename(cert_id)}_{sanitize_filename(last_name)}_{sanitize_filename(first_name)}.pdf"
                            return fname, pdf_bytes

                        tasks.append(render_one())
                    except Exception as e:
                        logger.error(f"Error preparing row {row_num}: {str(e)}")
                        if state:
                            state.errors += 1
                            state.message = f"Ошибка в строке {row_num}"
                            await emit(job_id)
                        continue

                async for fut in _as_completed_iter(tasks):
                    fname, pdf_bytes = await fut
                    zf.writestr(fname, pdf_bytes)
                    processed_count += 1
                    if state:
                        state.processed = processed_count
                        state.message = f"Готово {processed_count} из {total}"
                        await emit(job_id)

        if processed_count == 0:
            try:
                if not incoming_fields and txt:
                    reader2 = csv.DictReader(io.StringIO(txt), delimiter=delim)
                    incoming_fields = list(reader2.fieldnames or [])
            except Exception:
                incoming_fields = []
            required = [
                "Имя/Name", "Фамилия/Surname", "Название тренинга/Название",
                "Даты/Дата", "ID/Id", "(опц.) Город/City", "(опц.) Страна/Country"
            ]
            hint = (
                "CSV распознан, но ни одной корректной строки не найдено. "
                "Проверьте заголовки и обязательные поля. Требуемые колонки: "
                + ", ".join(required)
            )
            if state:
                state.stage = "error"
                state.message = "Нет валидных строк"
                await emit(job_id)
            return PlainTextResponse(hint + "\n" + f"Полученные колонки: {incoming_fields}", status_code=400)

        mem_zip.seek(0)
        zip_bytes = mem_zip.getvalue()
        if state:
            state.stage = "zipping"
            state.message = "Упаковка ZIP"
            await emit(job_id)

        if state:
            state.stage = "done"
            state.message = "Готово"
            await emit(job_id)

        return Response(
            content=zip_bytes,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=certificates.zip"},
        )

    except Exception as e:
        logger.error(f"Generation failed: {str(e)}")
        if job_id:
            state = get_progress(job_id)
            state.stage = "error"
            state.message = str(e)
            await emit(job_id)
        raise


@app.post("/generate-async")
async def generate_async(
    csv_file: UploadFile = File(...),
    mode: str = Form(...),                  # print | online
    job_id: Optional[str] = Form(None),
):
    try:
        logger.info(f"Starting ASYNC certificate generation for mode: {mode}")

        if not job_id:
            job_id = f"job-{int(time.time())}-{os.getpid()}-{id(csv_file)}"
        state = get_progress(job_id)
        state.stage = "uploading"
        state.message = "Загрузка файла"
        await emit(job_id)

        data = await csv_file.read()
        filename = (csv_file.filename or '').lower()
        rows_list: List[Dict[str, str]] = []

        if filename.endswith('.xlsx') or filename.endswith('.xlsm'):
            if not HAS_XLSX:
                raise RuntimeError('Поддержка Excel не установлена на сервере')
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            ws = wb.active
            headers: List[str] = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [(c if c is not None else '') for c in row]
                if not any(str(c).strip() for c in cells): continue
                if not headers:
                    candidate = [str(c).strip() for c in cells]
                    tech = all(h.lower().startswith('column') or h.lower().startswith('unnamed') or h == '' for h in candidate)
                    non_empty = [h for h in candidate if h]
                    if tech or len(non_empty) < 3: continue
                    headers = candidate; continue
                d: Dict[str, str] = {}
                for idx, h in enumerate(headers):
                    val = '' if idx >= len(cells) or cells[idx] is None else str(cells[idx])
                    d[h] = val
                rows_list.append(d)
            if not headers or not rows_list:
                raise ValueError('Excel: нераспознан заголовок или нет данных')
        else:
            raw_txt = data.decode('utf-8-sig', errors='ignore')
            txt, delim = normalize_csv_and_get_delimiter(raw_txt)
            if not txt:
                raise ValueError('CSV пустой или нераспознанный формат')
            dict_reader = csv.DictReader(io.StringIO(txt), delimiter=delim)
            rows_list = [row for row in dict_reader]

        total = len(rows_list)
        state.total = total
        state.processed = 0
        state.stage = "processing"
        state.message = "Обработка строк"
        await emit(job_id)

        async def worker():
            try:
                mem_zip = io.BytesIO()
                processed_count = 0
                loop = asyncio.get_event_loop()
                tasks = []
                with ThreadPoolExecutor(max_workers=2) as executor:
                    with zipfile.ZipFile(mem_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                        for row_num, row in enumerate(rows_list, 1):
                            try:
                                is_online = (mode == "online")
                                course     = _get_field(row, "course")
                                dates_raw  = _get_field(row, "dates")
                                first_name = _get_field(row, "first_name")
                                last_name  = _get_field(row, "last_name")
                                cert_id    = _get_field(row, "id")
                                city       = _get_field(row, "city")
                                country    = _get_field(row, "country")

                                if not (course and dates_raw and first_name and last_name and cert_id):
                                    logger.warning(f"Skipping row {row_num}: missing required fields")
                                    continue

                                parsed = parse_dates(dates_raw)
                                kind = pick_kind(parsed)
                                use_small = need_small_variant(f"{first_name} {last_name}")
                                variant = "small" if use_small else "normal"

                                group = "online" if is_online else "print"
                                docx_name = DOCX_MAP[group][kind][variant]
                                docx_path = os.path.join(TEMPLATES_DIR, docx_name)
                                if not os.path.exists(docx_path):
                                    raise FileNotFoundError(f"Template not found: {docx_path}")

                                context = format_dates_for_jinja(parsed)
                                context.update({
                                    "Имя": first_name,
                                    "Фамилия": last_name,
                                    "Тренинг": course,
                                    "Идентификатор": cert_id,
                                    "Город": city or context.get("Город", "Москва"),
                                    "Страна": country,
                                })

                                if group == "online":
                                    pad_name = NBSP * ONLINE_NAME_PAD_NBSP
                                    pad_course = NBSP * ONLINE_COURSE_PAD_NBSP
                                    context["Имя"] = pad_name + context.get("Имя", "")
                                    context["Тренинг"] = pad_course + context.get("Тренинг", "")

                                async def render_one(docx_path=docx_path, context=context, cert_id=cert_id,
                                                     last_name=last_name, first_name=first_name, group=group):
                                    adjust = (group == "online")
                                    pdf_bytes = await loop.run_in_executor(
                                        executor, render_docx_template, docx_path, context, adjust
                                    )
                                    fname = f"{sanitize_filename(cert_id)}_{sanitize_filename(last_name)}_{sanitize_filename(first_name)}.pdf"
                                    return fname, pdf_bytes

                                tasks.append(render_one())
                            except Exception as e:
                                logger.error(f"Error preparing row {row_num}: {str(e)}")
                                state.errors += 1
                                state.message = f"Ошибка в строке {row_num}"
                                await emit(job_id)
                                continue

                        async for fut in _as_completed_iter(tasks):
                            fname, pdf_bytes = await fut
                            zf.writestr(fname, pdf_bytes)
                            processed_count += 1
                            state.processed = processed_count
                            state.message = f"Готово {processed_count} из {total}"
                            await emit(job_id)

                if processed_count == 0:
                    state.stage = "error"
                    state.message = "Нет валидных строк"
                    await emit(job_id)
                    return

                mem_zip.seek(0)
                JOB_RESULTS[job_id] = mem_zip.getvalue()
                state.stage = "zipping"
                state.message = "Упаковка ZIP"
                await emit(job_id)

                state.stage = "done"
                state.message = "Готово"
                await emit(job_id)
            except Exception as e:
                logger.error(f"ASYNC Generation failed: {str(e)}")
                state.stage = "error"
                state.message = str(e)
                await emit(job_id)

        asyncio.create_task(worker())
        return {"job_id": job_id}

    except Exception as e:
        logger.error(f"generate-async init failed: {str(e)}")
        if job_id:
            state = get_progress(job_id)
            state.stage = "error"
            state.message = str(e)
            await emit(job_id)
        return PlainTextResponse(str(e), status_code=400)


@app.get("/download/{job_id}")
def download_result(job_id: str):
    data = JOB_RESULTS.get(job_id)
    if not data:
        return PlainTextResponse("Результат не готов или истёк", status_code=404)
    try:
        del JOB_RESULTS[job_id]
    except Exception:
        pass
    return Response(
        content=data,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=certificates.zip"},
    )
